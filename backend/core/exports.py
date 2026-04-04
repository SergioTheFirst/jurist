"""Case export helpers for DOCX and PDF reports."""

from __future__ import annotations

from io import BytesIO
import os
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT_RGB = RGBColor(0xB8, 0x96, 0x5A)
TEXT_RGB = RGBColor(0x22, 0x24, 0x2C)


class ReportExportService:
    """Create branded DOCX and PDF reports from archived cases."""

    def export_docx(self, case_data: dict[str, Any]) -> bytes:
        """Build a DOCX report and return its bytes."""

        document = Document()
        self._set_doc_margins(document)
        self._build_docx_header(document, case_data)
        self._build_docx_body(document, case_data)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def export_pdf(self, case_data: dict[str, Any]) -> bytes:
        """Build a simple branded PDF report and return its bytes."""

        document = fitz.open()
        page = document.new_page(width=595, height=842)
        font_path = self._find_unicode_font()

        header_rect = fitz.Rect(36, 28, 559, 84)
        page.draw_rect(header_rect, color=(0.72, 0.59, 0.35), fill=(0.09, 0.1, 0.14), width=1)
        self._insert_text(
            page,
            fitz.Point(48, 54),
            "ЮристАИ · LegalDesk",
            font_size=18,
            font_path=font_path,
            color=(0.83, 0.67, 0.43),
        )
        self._insert_text(
            page,
            fitz.Point(48, 72),
            "Юридическое заключение",
            font_size=10,
            font_path=font_path,
            color=(0.86, 0.85, 0.8),
        )

        body_lines = self._pdf_lines(case_data)
        body_rect = fitz.Rect(42, 104, 553, 800)
        page.insert_textbox(
            body_rect,
            "\n".join(body_lines),
            fontsize=10.5,
            fontfile=font_path,
            color=(0.87, 0.85, 0.8),
            lineheight=1.35,
        )

        return document.tobytes()

    @staticmethod
    def _set_doc_margins(document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def _build_docx_header(self, document: Document, case_data: dict[str, Any]) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run("ЮристАИ · LegalDesk")
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = ACCENT_RGB

        subtitle = document.add_paragraph()
        subtitle_run = subtitle.add_run("Юридическое заключение")
        subtitle_run.font.name = "Arial"
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = TEXT_RGB

        meta = document.add_paragraph()
        meta_run = meta.add_run(
            f"Дело #{case_data.get('id') or '—'} · {case_data.get('filename') or 'Без имени'} · "
            f"{case_data.get('created_at') or '—'}"
        )
        meta_run.font.name = "Arial"
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = TEXT_RGB

        separator = document.add_paragraph()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "B8965A")
        border.append(bottom)
        separator._p.get_or_add_pPr().append(border)

    def _build_docx_body(self, document: Document, case_data: dict[str, Any]) -> None:
        self._docx_section(document, "Краткое резюме", case_data.get("legal_summary") or case_data.get("summary") or "—")
        self._docx_section(document, "Рекомендации", case_data.get("recommendations") or "—")
        self._docx_section(document, "Анонимизированный текст", case_data.get("anonymized_text") or "—")

        document.add_paragraph()
        laws_heading = document.add_paragraph()
        laws_heading_run = laws_heading.add_run("Нормы законодательства")
        laws_heading_run.bold = True
        laws_heading_run.font.name = "Times New Roman"
        laws_heading_run.font.size = Pt(13)
        laws_heading_run.font.color.rgb = ACCENT_RGB
        self._docx_bullets(document, case_data.get("relevant_laws") or [], ("article", "title", "text"))

        practice_heading = document.add_paragraph()
        practice_heading_run = practice_heading.add_run("Судебная практика")
        practice_heading_run.bold = True
        practice_heading_run.font.name = "Times New Roman"
        practice_heading_run.font.size = Pt(13)
        practice_heading_run.font.color.rgb = ACCENT_RGB
        self._docx_bullets(document, case_data.get("court_practice") or [], ("case", "court", "outcome"))

    def _docx_section(self, document: Document, title: str, body: str) -> None:
        heading = document.add_paragraph()
        heading_run = heading.add_run(title)
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = ACCENT_RGB

        paragraph = document.add_paragraph()
        run = paragraph.add_run(body)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT_RGB

    @staticmethod
    def _docx_bullets(document: Document, items: list[dict[str, Any]], fields: tuple[str, ...]) -> None:
        if not items:
            document.add_paragraph("Нет данных.")
            return

        for item in items:
            values = [str(item.get(field) or "").strip() for field in fields]
            line = " | ".join(value for value in values if value)
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line or "Без деталей")
            run.font.name = "Arial"
            run.font.size = Pt(10)

    @staticmethod
    def _pdf_lines(case_data: dict[str, Any]) -> list[str]:
        lines = [
            f"Дело: #{case_data.get('id') or '—'}",
            f"Файл: {case_data.get('filename') or 'Без имени'}",
            f"Дата: {case_data.get('created_at') or '—'}",
            f"Источник анализа: {case_data.get('source') or '—'}",
            "",
            "КРАТКОЕ РЕЗЮМЕ",
            str(case_data.get("legal_summary") or case_data.get("summary") or "—"),
            "",
            "РЕКОМЕНДАЦИИ",
            str(case_data.get("recommendations") or "—"),
            "",
            "НОРМЫ ЗАКОНОДАТЕЛЬСТВА",
        ]

        laws = case_data.get("relevant_laws") or []
        if laws:
            lines.extend(
                f"• {item.get('article') or '—'} | {item.get('title') or '—'} | {item.get('text') or ''}"
                for item in laws
            )
        else:
            lines.append("• Нет данных")

        lines.extend(["", "СУДЕБНАЯ ПРАКТИКА"])
        practice = case_data.get("court_practice") or []
        if practice:
            lines.extend(
                f"• {item.get('case') or '—'} | {item.get('court') or '—'} | {item.get('outcome') or ''}"
                for item in practice
            )
        else:
            lines.append("• Нет данных")

        lines.extend(["", "АНОНИМИЗИРОВАННЫЙ ТЕКСТ", str(case_data.get("anonymized_text") or "—")])
        return lines

    @staticmethod
    def _find_unicode_font() -> str | None:
        candidates = [
            Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "arial.ttf",
            Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / "times.ttf",
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        return None

    @staticmethod
    def _insert_text(
        page: fitz.Page,
        point: fitz.Point,
        text: str,
        *,
        font_size: float,
        font_path: str | None,
        color: tuple[float, float, float],
    ) -> None:
        page.insert_text(
            point,
            text,
            fontsize=font_size,
            fontfile=font_path,
            color=color,
        )


report_export_service = ReportExportService()
