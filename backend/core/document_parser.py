"""Document parsing strategies for LegalDesk."""

from __future__ import annotations

from abc import ABC, abstractmethod
from io import BytesIO
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
import zipfile
from xml.etree import ElementTree


MAX_DOCUMENT_LENGTH = 500_000
SCANNED_THRESHOLD = 50
SUPPORTED_FILE_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".odt"}
DOCX_XML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
TRUNCATION_WARNING = "[ПРЕДУПРЕЖДЕНИЕ: Документ обрезан до 500 000 символов]"
SCANNED_PDF_MESSAGE = (
    "Документ является сканом или содержит слабый текстовый слой. OCR недоступен или не настроен.\n"
    "Для обработки таких PDF установите Tesseract OCR и пакет pytesseract, затем добавьте tesseract в PATH."
)
OCR_FAILED_PDF_MESSAGE = (
    "Документ является сканом или содержит слабый текстовый слой. Не удалось выполнить OCR.\n"
    "Проверьте установку Tesseract OCR и наличие языковых данных rus/eng."
)
ENCRYPTED_PDF_MESSAGE = "PDF защищён паролем. Снимите защиту перед загрузкой."
DOC_LIBREOFFICE_MESSAGE = (
    "Формат .doc (старый Word) требует LibreOffice или совместимый конвертер.\n"
    "Установите LibreOffice и добавьте soffice/libreoffice в PATH.\n"
    "Или сохраните документ в формате .docx"
)
OCR_LANGUAGES = os.getenv("LEGALDESK_OCR_LANGS", "rus+eng")
OCR_DPI = int(os.getenv("LEGALDESK_OCR_DPI", "300"))
OCR_TESSERACT_ENV = "TESSERACT_CMD"


class ParseError(RuntimeError):
    """Raised when document content cannot be extracted safely."""


class ParserStrategy(ABC):
    """Single-file parser strategy."""

    @abstractmethod
    def parse(self, path: Path) -> str:
        """Extract text from a document."""


class PdfParser(ParserStrategy):
    """Extract text from PDF using pdfplumber and PyMuPDF fallback."""

    def parse(self, path: Path) -> str:
        text = self._try_pdfplumber(path)
        if self._needs_ocr(text):
            fallback = self._try_pymupdf(path)
            text = self._pick_best_text(text, fallback)
        if self._needs_ocr(text):
            ocr_text = self._try_ocr(path)
            text = self._pick_best_text(text, ocr_text)
        if self._is_scanned(text):
            if self._ocr_is_available():
                raise ParseError(OCR_FAILED_PDF_MESSAGE)
            raise ParseError(SCANNED_PDF_MESSAGE)
        return text

    def _try_pdfplumber(self, path: Path) -> str:
        try:
            import pdfplumber  # type: ignore[import-not-found]
        except ModuleNotFoundError as exc:  # pragma: no cover - dependency error
            raise ParseError("Не установлен pdfplumber для чтения PDF") from exc

        try:
            pages: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    fragments: list[str] = []
                    page_text = self._pick_best_text(
                        page.extract_text() or "",
                        page.extract_text(layout=True) or "",
                    )
                    if page_text.strip():
                        fragments.append(page_text.strip())
                    for table in page.extract_tables() or []:
                        table_text = self._table_to_text(table)
                        if table_text:
                            fragments.append(table_text)
                    if fragments:
                        pages.append("\n\n".join(fragments))
            return "\n\n".join(pages)
        except Exception as exc:  # pragma: no cover - depends on parser internals
            if self._looks_like_encrypted_error(exc):
                raise ParseError(ENCRYPTED_PDF_MESSAGE) from exc
            return ""

    def _try_pymupdf(self, path: Path) -> str:
        try:
            import fitz  # type: ignore[import-not-found]
        except ModuleNotFoundError as exc:  # pragma: no cover - dependency error
            raise ParseError("Не установлен PyMuPDF для резервного чтения PDF") from exc

        try:
            pages: list[str] = []
            with fitz.open(path) as doc:
                if getattr(doc, "needs_pass", False):
                    raise ParseError(ENCRYPTED_PDF_MESSAGE)
                for page in doc:
                    page_text = self._extract_sorted_page_text(page) or page.get_text("text", sort=True) or page.get_text("text") or ""
                    if page_text.strip():
                        pages.append(page_text.strip())
            return "\n\n".join(pages)
        except ParseError:
            raise
        except Exception as exc:  # pragma: no cover - depends on parser internals
            if self._looks_like_encrypted_error(exc):
                raise ParseError(ENCRYPTED_PDF_MESSAGE) from exc
            return ""

    @staticmethod
    def _table_to_text(table: list[list[str | None]] | None) -> str:
        if not table:
            return ""
        lines: list[str] = []
        for row in table:
            cleaned = [PdfParser._clean_cell(cell) for cell in row]
            if any(cell for cell in cleaned):
                lines.append("| " + " | ".join(cleaned) + " |")
        return "\n".join(lines)

    @staticmethod
    def _clean_cell(value: str | None) -> str:
        return " ".join((value or "").split())

    @staticmethod
    def _is_scanned(text: str) -> bool:
        return len(text.strip()) < SCANNED_THRESHOLD

    @classmethod
    def _needs_ocr(cls, text: str) -> bool:
        return cls._is_scanned(text) or cls._looks_corrupted(text)

    @staticmethod
    def _looks_like_encrypted_error(exc: Exception) -> bool:
        message = str(exc).lower()
        return "password" in message or "encrypted" in message

    @staticmethod
    def _extract_sorted_page_text(page: object) -> str:
        try:
            blocks = page.get_text("blocks")  # type: ignore[attr-defined]
        except Exception:  # pragma: no cover - backend-specific
            return ""

        sortable_blocks: list[tuple[float, float, str]] = []
        for block in blocks or []:
            if len(block) < 5:
                continue
            x0, y0, _x1, _y1, text, *_rest = block
            cleaned = str(text).strip()
            if cleaned:
                sortable_blocks.append((float(x0), float(y0), cleaned))
        return "\n\n".join(
            text for _x, _y, text in sorted(sortable_blocks, key=lambda item: (round(item[0], 1), round(item[1], 1)))
        )

    @staticmethod
    def _pick_best_text(*candidates: str) -> str:
        non_empty = [candidate for candidate in candidates if candidate and candidate.strip()]
        if not non_empty:
            return ""
        return max(non_empty, key=PdfParser._text_quality_score)

    def _try_ocr(self, path: Path) -> str:
        if not self._ocr_is_available():
            return ""

        for extractor in (self._ocr_with_pymupdf, self._ocr_with_pytesseract):
            try:
                text = extractor(path)
            except ParseError:
                raise
            except Exception:
                continue
            if text.strip():
                return text
        return ""

    @staticmethod
    def _resolve_tesseract_command() -> str | None:
        configured = os.getenv(OCR_TESSERACT_ENV, "").strip()
        if configured:
            return configured
        return shutil.which("tesseract")

    def _ocr_is_available(self) -> bool:
        return bool(self._resolve_tesseract_command())

    def _ocr_with_pymupdf(self, path: Path) -> str:
        try:
            import fitz  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            return ""

        if not hasattr(fitz.Page, "get_textpage_ocr"):
            return ""

        tesseract_cmd = self._resolve_tesseract_command()
        if not tesseract_cmd:
            return ""

        os.environ.setdefault("TESSERACT", tesseract_cmd)
        pages: list[str] = []
        with fitz.open(path) as doc:
            if getattr(doc, "needs_pass", False):
                raise ParseError(ENCRYPTED_PDF_MESSAGE)
            for page in doc:
                text_page = page.get_textpage_ocr(
                    language=OCR_LANGUAGES,
                    dpi=OCR_DPI,
                    full=True,
                    tessdata=os.getenv("TESSDATA_PREFIX") or None,
                )
                page_text = page.get_text("text", textpage=text_page, sort=True) or ""
                if page_text.strip():
                    pages.append(page_text.strip())
        return "\n\n".join(pages)

    def _ocr_with_pytesseract(self, path: Path) -> str:
        try:
            import fitz  # type: ignore[import-not-found]
            from PIL import Image, ImageOps  # type: ignore[import-not-found]
            import pytesseract  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            return ""

        tesseract_cmd = self._resolve_tesseract_command()
        if not tesseract_cmd:
            return ""

        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        scale = max(OCR_DPI / 72.0, 1.0)
        matrix = fitz.Matrix(scale, scale)
        pages: list[str] = []
        with fitz.open(path) as doc:
            if getattr(doc, "needs_pass", False):
                raise ParseError(ENCRYPTED_PDF_MESSAGE)
            for page in doc:
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                prepared = ImageOps.autocontrast(image.convert("L")).point(lambda value: 255 if value > 180 else 0)
                page_text = pytesseract.image_to_string(
                    prepared,
                    lang=OCR_LANGUAGES,
                    config="--oem 3 --psm 6 -c preserve_interword_spaces=1",
                )
                if page_text.strip():
                    pages.append(page_text.strip())
        return "\n\n".join(pages)

    @staticmethod
    def _looks_corrupted(text: str) -> bool:
        stripped = text.strip()
        if not stripped:
            return False
        question_ratio = stripped.count("?") / max(len(stripped), 1)
        replacement_ratio = stripped.count("\ufffd") / max(len(stripped), 1)
        single_letter_runs = len(re.findall(r"(?:\b[А-Яа-яA-Za-zЁё]\b(?:\s+|$)){5,}", stripped))
        return question_ratio > 0.12 or replacement_ratio > 0.02 or single_letter_runs > 0

    @staticmethod
    def _text_quality_score(text: str) -> float:
        stripped = text.strip()
        if not stripped:
            return float("-inf")

        letters = sum(1 for char in stripped if char.isalpha())
        digits = sum(1 for char in stripped if char.isdigit())
        spaces = stripped.count(" ")
        bad = stripped.count("?") + stripped.count("\ufffd") * 8
        single_letter_runs = len(re.findall(r"(?:\b[А-Яа-яA-Za-zЁё]\b(?:\s+|$)){5,}", stripped))
        return float((letters * 2) + digits + min(spaces, 120) - (bad * 6) - (single_letter_runs * 30))


class DocxParser(ParserStrategy):
    """Extract plain text from DOCX paragraphs, tables, headers, and text boxes."""

    def parse(self, path: Path) -> str:
        text = self._parse_with_python_docx(path)
        if text.strip():
            return text
        return self._parse_with_xml_fallback(path)

    def _parse_with_python_docx(self, path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except ModuleNotFoundError as exc:  # pragma: no cover - dependency error
            raise ParseError("Не установлен python-docx для чтения DOCX") from exc

        try:
            document = Document(path)
            chunks: list[str] = []

            for section in document.sections:
                header_text = self._collect_paragraphs(section.header.paragraphs)
                if header_text:
                    chunks.append(f"[ВЕРХНИЙ КОЛОНТИТУЛ: {header_text}]")
                footer_text = self._collect_paragraphs(section.footer.paragraphs)
                if footer_text:
                    chunks.append(f"[НИЖНИЙ КОЛОНТИТУЛ: {footer_text}]")

            paragraph_text = self._collect_paragraphs(document.paragraphs)
            if paragraph_text:
                chunks.append(paragraph_text)

            for table in document.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [" ".join(cell.text.split()) for cell in row.cells]
                    if any(cells):
                        rows.append("| " + " | ".join(cells) + " |")
                if rows:
                    chunks.append("\n".join(rows))

            text_box_text = self._extract_text_boxes(document)
            if text_box_text:
                chunks.append(text_box_text)

            return "\n\n".join(chunk for chunk in chunks if chunk.strip())
        except Exception:
            return ""

    @staticmethod
    def _collect_paragraphs(paragraphs: list[object]) -> str:
        lines = [
            " ".join(getattr(paragraph, "text", "").split())
            for paragraph in paragraphs
            if getattr(paragraph, "text", "").strip()
        ]
        return "\n".join(lines)

    @staticmethod
    def _extract_text_boxes(document: object) -> str:
        try:
            elements = document.element.xpath(".//w:txbxContent//w:t")  # type: ignore[attr-defined]
        except Exception:  # pragma: no cover - depends on XML content
            return ""
        lines = [" ".join(getattr(element, "text", "").split()) for element in elements if getattr(element, "text", "").strip()]
        return "\n".join(lines)

    def _parse_with_xml_fallback(self, path: Path) -> str:
        try:
            with zipfile.ZipFile(path) as archive:
                chunks: list[str] = []
                chunks.extend(self._extract_xml_parts(archive, "word/header", "ВЕРХНИЙ КОЛОНТИТУЛ"))
                body_text = self._extract_document_xml(archive, "word/document.xml")
                if body_text:
                    chunks.append(body_text)
                chunks.extend(self._extract_xml_parts(archive, "word/footer", "НИЖНИЙ КОЛОНТИТУЛ"))
                return "\n\n".join(chunk for chunk in chunks if chunk.strip())
        except Exception as exc:
            raise ParseError(f"Не удалось прочитать DOCX: {exc}") from exc

    def _extract_xml_parts(self, archive: zipfile.ZipFile, prefix: str, label: str) -> list[str]:
        chunks: list[str] = []
        for member in sorted(name for name in archive.namelist() if name.startswith(prefix) and name.endswith(".xml")):
            text = self._extract_document_xml(archive, member)
            if text:
                chunks.append(f"[{label}: {text}]")
        return chunks

    def _extract_document_xml(self, archive: zipfile.ZipFile, member: str) -> str:
        try:
            xml_bytes = archive.read(member)
        except KeyError:
            return ""

        root = ElementTree.fromstring(xml_bytes)
        parent = root.find("w:body", DOCX_XML_NS) if root.tag.endswith("document") else root
        if parent is None:
            return ""

        chunks: list[str] = []
        for child in list(parent):
            if child.tag.endswith("}p"):
                paragraph_text = self._paragraph_from_xml(child)
                if paragraph_text:
                    chunks.append(paragraph_text)
            elif child.tag.endswith("}tbl"):
                table_text = self._table_from_xml(child)
                if table_text:
                    chunks.append(table_text)
        return "\n\n".join(chunk for chunk in chunks if chunk.strip())

    @staticmethod
    def _paragraph_from_xml(node: ElementTree.Element) -> str:
        parts = [" ".join((element.text or "").split()) for element in node.findall(".//w:t", DOCX_XML_NS)]
        cleaned = [part for part in parts if part]
        return " ".join(cleaned).strip()

    @staticmethod
    def _table_from_xml(node: ElementTree.Element) -> str:
        rows: list[str] = []
        for row in node.findall("./w:tr", DOCX_XML_NS):
            cells: list[str] = []
            for cell in row.findall("./w:tc", DOCX_XML_NS):
                paragraphs = [DocxParser._paragraph_from_xml(paragraph) for paragraph in cell.findall("./w:p", DOCX_XML_NS)]
                cell_text = " ".join(part for part in paragraphs if part).strip()
                cells.append(cell_text)
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")
        return "\n".join(rows)


class DocParser(ParserStrategy):
    """Convert legacy DOC to DOCX via LibreOffice and reuse the DOCX parser."""

    def __init__(self, docx_parser: DocxParser) -> None:
        self._docx_parser = docx_parser

    def parse(self, path: Path) -> str:
        executable = shutil.which("libreoffice") or shutil.which("soffice")
        if executable is None:
            raise ParseError(DOC_LIBREOFFICE_MESSAGE)

        with tempfile.TemporaryDirectory(prefix="legaldesk-doc-") as tmp_dir_str:
            tmp_dir = Path(tmp_dir_str)
            command = [
                executable,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_dir),
                str(path),
            ]
            try:
                subprocess.run(
                    command,
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=10,
                )
            except subprocess.TimeoutExpired as exc:
                raise ParseError("Превышено время конвертации .doc через LibreOffice (10 секунд).") from exc
            except FileNotFoundError as exc:
                raise ParseError(DOC_LIBREOFFICE_MESSAGE) from exc
            except subprocess.CalledProcessError as exc:
                raise ParseError(f"LibreOffice не смог конвертировать .doc: {exc.stderr.strip() or exc.stdout.strip()}") from exc

            converted = tmp_dir / f"{path.stem}.docx"
            if not converted.exists():
                raise ParseError("LibreOffice не создал .docx после конвертации .doc")
            return self._docx_parser.parse(converted)


class TxtParser(ParserStrategy):
    """Read text files with several encoding fallbacks."""

    def parse(self, path: Path) -> str:
        raw_bytes = path.read_bytes()
        for encoding in ("utf-8", "cp1251", "cp866"):
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        detected_encoding: str | None = None
        try:
            import chardet  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            detected_encoding = None
        else:
            detected = chardet.detect(raw_bytes)
            detected_encoding = detected.get("encoding")

        if detected_encoding:
            try:
                return raw_bytes.decode(detected_encoding)
            except UnicodeDecodeError:
                pass

        return raw_bytes.decode("latin-1")


class RtfParser(ParserStrategy):
    """Extract plain text from RTF documents."""

    def parse(self, path: Path) -> str:
        raw_text = TxtParser().parse(path)
        try:
            from striprtf.striprtf import rtf_to_text  # type: ignore[import-not-found]
        except ModuleNotFoundError:
            return self._regex_fallback(raw_text)
        return rtf_to_text(raw_text)

    @staticmethod
    def _regex_fallback(text: str) -> str:
        without_groups = re.sub(r"{\\[^{}]*}", " ", text)
        without_controls = re.sub(r"\\[a-zA-Z]+\d* ?", " ", without_groups)
        without_braces = re.sub(r"[{}]", " ", without_controls)
        return re.sub(r"\s+", " ", without_braces)


class OdtParser(ParserStrategy):
    """Extract text paragraphs from ODT documents."""

    def parse(self, path: Path) -> str:
        try:
            from odf import text as odf_text  # type: ignore[import-not-found]
            from odf.opendocument import load as odf_load  # type: ignore[import-not-found]
            from odf.teletype import extractText  # type: ignore[import-not-found]
        except ModuleNotFoundError as exc:  # pragma: no cover - dependency error
            raise ParseError("Не установлен odfpy для чтения ODT") from exc

        document = odf_load(str(path))
        lines = [
            " ".join(extractText(node).split())
            for node in document.getElementsByType(odf_text.P)
            if extractText(node).strip()
        ]
        return "\n".join(lines)


class DocumentExtractor:
    """Facade for format-specific parsing strategies."""

    def __init__(self) -> None:
        docx_parser = DocxParser()
        self._parsers: dict[str, ParserStrategy] = {
            ".pdf": PdfParser(),
            ".docx": docx_parser,
            ".doc": DocParser(docx_parser),
            ".txt": TxtParser(),
            ".rtf": RtfParser(),
            ".odt": OdtParser(),
        }

    def extract(self, path: str | Path) -> str:
        file_path = Path(path)
        suffix = file_path.suffix.lower()
        parser = self._parsers.get(suffix)
        if parser is None:
            raise ParseError(f"Формат '{suffix}' не поддерживается")

        extracted = parser.parse(file_path)
        min_length = SCANNED_THRESHOLD if suffix == ".pdf" else 1
        return self._postprocess(extracted, min_length=min_length)

    def extract_from_text(self, text: str) -> str:
        if not text or not text.strip():
            raise ParseError("Текст не может быть пустым")
        normalized = self._normalize_text(text)
        if len(normalized) > MAX_DOCUMENT_LENGTH:
            raise ParseError("Текст слишком длинный для анализа (максимум 500 000 символов)")
        if len(normalized) < 20:
            raise ParseError("Текст слишком короткий для анализа (минимум 20 символов)")
        return normalized

    @staticmethod
    def _postprocess(text: str, *, min_length: int) -> str:
        normalized = DocumentExtractor._normalize_text(text)
        if len(normalized) < min_length:
            raise ParseError("Документ пуст")
        if len(normalized) > MAX_DOCUMENT_LENGTH:
            truncated_body = normalized[:MAX_DOCUMENT_LENGTH]
            normalized = f"{TRUNCATION_WARNING}\n\n{truncated_body}".strip()
        return normalized

    @staticmethod
    def _normalize_text(text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = normalized.lstrip("\ufeff")
        normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()


extractor = DocumentExtractor()
