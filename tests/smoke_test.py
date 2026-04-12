"""Runtime smoke tests for the LegalDesk processing pipeline."""

from __future__ import annotations

from collections.abc import Callable
import io
import json
import threading
import time
from pathlib import Path
from types import SimpleNamespace

import httpx
import pytest
from fastapi.testclient import TestClient


@pytest.fixture()
def client(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    """Provide an isolated FastAPI client backed by a temporary archive DB."""

    monkeypatch.setenv("KP_API_KEY", "DEMO_KEY")
    monkeypatch.delenv("LLM_BASE_URL", raising=False)
    monkeypatch.delenv("LLM_MODEL", raising=False)
    monkeypatch.setenv("ANALYSIS_MODE", "auto")

    from backend.core import audit as audit_module
    from backend.core import archive as archive_module
    from backend.core.audit import AuditLogManager
    from backend.core.archive import ArchiveManager
    import backend.main as main_module

    temp_root = tmp_path
    uploads_dir = temp_root / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    temp_archive = ArchiveManager(temp_root / "cases.db")
    temp_audit = AuditLogManager(temp_root / "audit.db")

    monkeypatch.setattr(main_module, "archive", temp_archive)
    monkeypatch.setattr(main_module, "audit_log", temp_audit)
    monkeypatch.setattr(archive_module, "archive", temp_archive)
    monkeypatch.setattr(audit_module, "audit_log", temp_audit)
    monkeypatch.setattr(main_module, "UPLOAD_DIR", uploads_dir)

    with TestClient(main_module.app) as test_client:
        yield test_client


def test_anonymizer_fio() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Иванов Иван Иванович подал заявление")
    assert "[ФИО1]" in result.anonymized_text
    assert result.total_replacements >= 1


def test_anonymizer_inn() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("ИНН: 771234567890")
    assert "[ИНН1]" in result.anonymized_text


def test_anonymizer_phone() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Телефон: +7 (999) 123-45-67")
    assert "[ТЕЛЕФОН1]" in result.anonymized_text


def test_anonymizer_email() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Email: ivanov@mail.ru")
    assert "[EMAIL1]" in result.anonymized_text


def test_anonymizer_bank_account() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Р/с 40817810099910004312, БИК 044525225")
    assert "[БАНК. СЧЁТ1]" in result.anonymized_text
    assert "[БИК1]" in result.anonymized_text


def test_anonymizer_vehicle_identifiers() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Госномер А123ВС777, VIN 1HGBH41JXMN109186, СТС 77 МА 123456")
    assert "[ГОСНОМЕР1]" in result.anonymized_text
    assert "[VIN1]" in result.anonymized_text
    assert "[СТС1]" in result.anonymized_text


def test_anonymizer_insurance_policy() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("Полис ТТТ 1234567890 предоставлен в материалы дела")
    assert "[СТРАХ. ПОЛИС1]" in result.anonymized_text


@pytest.mark.parametrize(
    "text",
    [
        "Заявитель проживает по адресу: 119019, г. Москва, ул. Арбат, д. 12, кв. 5.",
        "Фактический адрес: Рязанская обл., Клепиковский р-н, дер. Ласково, д. 7.",
        "Зарегистрирован по адресу: г. Волгоград, ул. им. Гагарина, д. 4.",
        "Истец проживает: г. Москва, ул. 1-я Тверская-Ямская, д. 11, корп. 2.",
    ],
    ids=["full_address", "rural_address", "im_prefix", "compound_street"],
)
def test_anonymizer_address(text: str) -> None:
    """Address candidates must be detected and replaced with [АДРЕС1]."""

    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize(text)
    assert "[АДРЕС1]" in result.anonymized_text, (
        f"Expected [АДРЕС1] in anonymized text, got: {result.anonymized_text}"
    )
    addr_replacements = [r for r in result.replacements if r.entity_type in ("АДРЕС", "LOC")]
    assert len(addr_replacements) >= 1, (
        f"Expected at least one АДРЕС replacement, got: {result.replacements}"
    )


def test_numbered_placeholders_distinguish_and_reuse_same_type_entities() -> None:
    from backend.core.anonymizer import _CandidateSpan, anonymizer

    text = "ООО Ромашка и ООО Ландыш заключили договор с ООО Ромашка."
    first_org = "ООО Ромашка"
    second_org = "ООО Ландыш"
    first_start = text.index(first_org)
    second_start = text.index(second_org)
    repeated_start = text.rindex(first_org)
    spans = [
        _CandidateSpan(
            id="org-1",
            original=first_org,
            placeholder="[ОРГАНИЗАЦИЯ]",
            entity_type="ORG",
            start=first_start,
            end=first_start + len(first_org),
            source="ner",
            confidence=0.95,
        ),
        _CandidateSpan(
            id="org-2",
            original=second_org,
            placeholder="[ОРГАНИЗАЦИЯ]",
            entity_type="ORG",
            start=second_start,
            end=second_start + len(second_org),
            source="ner",
            confidence=0.95,
        ),
        _CandidateSpan(
            id="org-3",
            original=first_org,
            placeholder="[ОРГАНИЗАЦИЯ]",
            entity_type="ORG",
            start=repeated_start,
            end=repeated_start + len(first_org),
            source="ner",
            confidence=0.95,
        ),
    ]

    anonymized_text, records, entities_found = anonymizer._render_replacements(text, spans)

    assert anonymized_text == "[ОРГАНИЗАЦИЯ1] и [ОРГАНИЗАЦИЯ2] заключили договор с [ОРГАНИЗАЦИЯ1]."
    assert [record.placeholder for record in records] == [
        "[ОРГАНИЗАЦИЯ1]",
        "[ОРГАНИЗАЦИЯ2]",
        "[ОРГАНИЗАЦИЯ1]",
    ]
    assert entities_found["ORG"] == ["ООО Ромашка", "ООО Ландыш", "ООО Ромашка"]


def test_numbered_placeholders_for_multiple_regex_entities() -> None:
    from backend.core.anonymizer import anonymizer

    text = "ИНН 771234567890 и ИНН 500100732259 указаны в договоре."
    result = anonymizer.anonymize(text)

    assert "[ИНН1]" in result.anonymized_text
    assert "[ИНН2]" in result.anonymized_text
    assert len([record for record in result.replacements if record.entity_type == "ИНН"]) == 2


def test_anonymizer_keeps_specific_org_names_even_with_bare_form_whitelist() -> None:
    from backend.core.anonymizer import anonymizer

    text = 'ООО "Ромашка" заключило договор с ООО "Ландыш".'
    result = anonymizer.anonymize(text)

    assert "[ОРГАНИЗАЦИЯ1]" in result.anonymized_text
    assert "[ОРГАНИЗАЦИЯ2]" in result.anonymized_text
    assert "ООО" not in result.whitelist_skipped


def test_person_rule_expands_surname_to_full_name() -> None:
    from backend.core.anonymizer import _CandidateSpan
    from backend.core.entity_rules.per_rules import PersonRuleEngine

    text = "Истец Иванов Иван Иванович подал иск."
    start = text.index("Иванов")
    candidate = _CandidateSpan(
        id="per-1",
        original="Иванов",
        placeholder="[ФИО]",
        entity_type="PER",
        start=start,
        end=start + len("Иванов"),
        source="ner",
        confidence=0.86,
    )

    expanded = PersonRuleEngine("[ФИО]").expand_candidate(text, candidate)
    assert expanded.original == "Иванов Иван Иванович"


def test_person_rule_rejects_generic_capitalized_legal_phrase() -> None:
    from backend.core.anonymizer import _CandidateSpan
    from backend.core.entity_rules.per_rules import PersonRuleEngine

    text = "Трудовой Договор заключен в письменной форме."
    start = text.index("Трудовой")
    candidate = _CandidateSpan(
        id="per-2",
        original="Трудовой Договор",
        placeholder="[ФИО]",
        entity_type="PER",
        start=start,
        end=start + len("Трудовой Договор"),
        source="ner",
        confidence=0.94,
    )

    decision = PersonRuleEngine("[ФИО]").validate_candidate(text, candidate)
    assert decision.accepted is False


def test_anonymizer_supplements_person_initials_patterns() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Представитель Иванов И.И. подписал договор и направил его в суд."
    result = anonymizer.anonymize(text)

    assert "[ФИО1]" in result.anonymized_text
    assert any(record.original == "Иванов И.И." for record in result.replacements)


def test_person_aliases_reuse_same_placeholder() -> None:
    from backend.core.anonymizer import anonymizer

    text = (
        "Иванов Иван Иванович подписал договор. "
        "Далее Иванов направил копию в суд. "
        "И.И. Иванов приложил доверенность."
    )
    result = anonymizer.anonymize(text)

    assert result.anonymized_text.count("[ФИО1]") == 3
    assert "[ФИО2]" not in result.anonymized_text


def test_person_alias_prefix_form_reuses_full_name_placeholder() -> None:
    from backend.core.anonymizer import anonymizer

    text = (
        "Петров Петр Петрович подал заявление. "
        "гр. Петров получил копию определения."
    )
    result = anonymizer.anonymize(text)

    assert result.anonymized_text.count("[ФИО1]") == 2
    assert "[ФИО2]" not in result.anonymized_text


def test_anonymizer_empty() -> None:
    from backend.core.anonymizer import anonymizer

    result = anonymizer.anonymize("")
    assert result.anonymized_text == ""
    assert len(result.warnings) > 0


def test_extract_from_text_valid() -> None:
    from backend.core.document_parser import extractor

    text = "Стороны договорились о расторжении договора аренды."
    result = extractor.extract_from_text(text)
    assert result == text.strip()


def test_extract_from_text_too_short() -> None:
    from backend.core.document_parser import ParseError, extractor

    with pytest.raises(ParseError):
        extractor.extract_from_text("кратко")


def test_extract_from_text_empty() -> None:
    from backend.core.document_parser import ParseError, extractor

    with pytest.raises(ParseError):
        extractor.extract_from_text("   ")


def test_parse_txt_utf8(tmp_path: pytest.TempPathFactory) -> None:
    from backend.core.document_parser import extractor

    file_path = tmp_path / "test.txt"
    file_path.write_text("Договор № 123 от 01.01.2024\nСторона: ООО Ромашка", encoding="utf-8")
    result = extractor.extract(file_path)
    assert "Договор" in result


def test_parse_txt_cp1251(tmp_path: pytest.TempPathFactory) -> None:
    from backend.core.document_parser import extractor

    file_path = tmp_path / "test_cp1251.txt"
    file_path.write_bytes("Иванов Иван Иванович".encode("cp1251"))
    result = extractor.extract(file_path)
    assert "Иванов" in result


def test_parse_docx(tmp_path: pytest.TempPathFactory) -> None:
    from docx import Document
    from backend.core.document_parser import extractor

    doc = Document()
    doc.add_paragraph("Настоящий договор заключён между сторонами.")
    doc.add_paragraph("Предмет договора: оказание юридических услуг.")
    path = tmp_path / "test.docx"
    doc.save(str(path))

    result = extractor.extract(path)
    assert "договор" in result.lower()
    assert len(result) > 20


def test_parse_docx_xml_fallback(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from docx import Document
    from backend.core.document_parser import DocxParser

    doc = Document()
    doc.add_paragraph("Договор поставки заключён между сторонами.")
    doc.add_paragraph("Телефон для связи: +7 495 000 00 00.")
    path = tmp_path / "fallback.docx"
    doc.save(str(path))

    monkeypatch.setattr(DocxParser, "_parse_with_python_docx", lambda self, _path: "")

    result = DocxParser().parse(path)
    assert "договор поставки" in result.lower()
    assert "+7 495 000 00 00" in result


def test_parse_doc_requires_converter(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from backend.core.document_parser import DocParser, DocxParser, ParseError

    path = tmp_path / "legacy.doc"
    path.write_bytes(b"legacy-doc")
    monkeypatch.setattr("backend.core.document_parser.shutil.which", lambda *_args, **_kwargs: None)

    with pytest.raises(ParseError, match="LibreOffice|конвертер"):
        DocParser(DocxParser()).parse(path)


def test_pdf_text_quality_prefers_readable_candidate() -> None:
    from backend.core.document_parser import PdfParser

    broken = "? ? ? ? ? ????? ??????"
    readable = "Договор аренды заключён между сторонами."

    assert PdfParser._pick_best_text(broken, readable) == readable


def test_pdf_scanned_uses_ocr_when_available(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from backend.core.document_parser import PdfParser

    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.4 scan")

    monkeypatch.setattr(PdfParser, "_try_pdfplumber", lambda self, _path: "")
    monkeypatch.setattr(PdfParser, "_try_pymupdf", lambda self, _path: "")
    monkeypatch.setattr(PdfParser, "_ocr_is_available", lambda self: True)
    monkeypatch.setattr(
        PdfParser,
        "_try_ocr",
        lambda self, _path: "Иванов Иван Иванович\nДоговор аренды заключён между сторонами.",
    )

    result = PdfParser().parse(path)
    assert "Иванов Иван Иванович" in result
    assert "Договор аренды" in result


def test_pdf_weak_text_prefers_ocr_result(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from backend.core.document_parser import PdfParser

    path = tmp_path / "weak.pdf"
    path.write_bytes(b"%PDF-1.4 weak")

    monkeypatch.setattr(PdfParser, "_try_pdfplumber", lambda self, _path: "? ? ? ? ? ????? ??????")
    monkeypatch.setattr(PdfParser, "_try_pymupdf", lambda self, _path: "? ? ? ??????")
    monkeypatch.setattr(PdfParser, "_ocr_is_available", lambda self: True)
    monkeypatch.setattr(
        PdfParser,
        "_try_ocr",
        lambda self, _path: "Договор поставки между ООО Ромашка и ИП Петров заключён 12.01.2026.",
    )

    result = PdfParser().parse(path)
    assert result == "Договор поставки между ООО Ромашка и ИП Петров заключён 12.01.2026."


def test_pdf_scanned_without_ocr_raises_parse_error(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from backend.core.document_parser import ParseError, PdfParser

    path = tmp_path / "scan-no-ocr.pdf"
    path.write_bytes(b"%PDF-1.4 scan")

    monkeypatch.setattr(PdfParser, "_try_pdfplumber", lambda self, _path: "")
    monkeypatch.setattr(PdfParser, "_try_pymupdf", lambda self, _path: "")
    monkeypatch.setattr(PdfParser, "_ocr_is_available", lambda self: False)

    with pytest.raises(ParseError, match="сканом|OCR"):
        PdfParser().parse(path)


def test_archive_save_and_retrieve(tmp_path: pytest.TempPathFactory) -> None:
    from backend.core.archive import ArchiveManager

    db_path = tmp_path / "test.db"
    archive_manager = ArchiveManager(db_path=db_path)
    legal = SimpleNamespace(
        summary="Тест",
        relevant_laws=[{"title": "ТК РФ", "article": "ст.81", "text": "..."}],
        court_practice=[],
        recommendations="Проверить документы",
    )

    case_id = archive_manager.save_case(
        filename="test.docx",
        anonymized_text="Текст без ПДн",
        entities_found={"PER": ["Иванов"]},
        legal_result=legal,
        total_replacements=1,
        input_type="file",
    )

    assert case_id is not None
    case = archive_manager.get_case(case_id)
    assert case is not None
    assert case["filename"] == "test.docx"
    assert case["input_type"] == "file"
    assert case["total_replacements"] == 1


def test_runtime_data_root_uses_env_override(tmp_path: pytest.TempPathFactory, monkeypatch: pytest.MonkeyPatch) -> None:
    from backend import runtime_paths

    target = tmp_path / "portable-data"
    monkeypatch.setenv("LEGALDESK_DATA_DIR", str(target))

    assert runtime_paths.runtime_data_root() == target
    assert runtime_paths.archive_db_path().parent == target / "archive"
    assert runtime_paths.audit_db_path().parent == target / "archive"


def test_desktop_launcher_builds_script_mode_server_command(
    tmp_path: pytest.TempPathFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from desktop import launcher

    monkeypatch.setattr(launcher.sys, "executable", r"C:\Python\python.exe")
    monkeypatch.setattr(launcher.sys, "frozen", False, raising=False)

    command = launcher.build_server_command("127.0.0.1", 8000, tmp_path)
    assert command[0].endswith("python.exe")
    assert command[1].endswith("launcher.py")
    assert "--serve" in command
    assert str(tmp_path) in command


def test_desktop_launcher_builds_frozen_server_command(
    tmp_path: pytest.TempPathFactory,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from desktop import launcher

    monkeypatch.setattr(launcher.sys, "executable", r"C:\Program Files\LegalDesk\LegalDesk.exe")
    monkeypatch.setattr(launcher.sys, "frozen", True, raising=False)

    command = launcher.build_server_command("127.0.0.1", 8000, tmp_path)
    assert command[0].endswith("LegalDesk.exe")
    assert not any(part.endswith("launcher.py") for part in command)
    assert "--serve" in command
    assert str(tmp_path) in command


def test_health(client: TestClient) -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert "kp_mode" in response.json()


def test_shutdown_endpoint_calls_local_handler(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    import backend.main as main_module

    called: dict[str, bool] = {"done": False}

    def fake_shutdown() -> None:
        called["done"] = True

    monkeypatch.setattr(main_module.app.state, "shutdown_handler", fake_shutdown, raising=False)

    response = client.post("/api/system/shutdown")

    assert response.status_code == 200
    assert response.json()["status"] == "stopping"
    assert called["done"] is True


def test_process_text_short_fails(client: TestClient) -> None:
    response = client.post("/api/process-text", json={"text": "кратко"})
    assert response.status_code == 422


def test_process_text_valid(client: TestClient) -> None:
    long_text = "Настоящий договор заключён между ООО Альфа и ИП Петров. " * 5
    response = client.post("/api/process-text", json={"text": long_text})
    assert response.status_code == 200
    body = response.json()
    assert "case_id" in body
    assert "anonymization" in body
    assert "legal_analysis" in body
    assert body["anonymization"]["anonymized_text"] != ""


def test_process_file_txt(client: TestClient) -> None:
    content = (
        "Договор аренды. Арендатор: Сидоров Сергей Петрович. "
        "Телефон: +7 (999) 000-11-22. ИНН: 500100732259. "
    ) * 3
    response = client.post(
        "/api/process",
        files={"file": ("doc.txt", io.BytesIO(content.encode("utf-8")), "text/plain")},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["anonymization"]["total_replacements"] >= 1


def test_preview_file_docx(client: TestClient, tmp_path: pytest.TempPathFactory) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Настоящий договор заключён между сторонами.")
    doc.add_paragraph("Предмет договора: оказание юридических услуг.")
    path = tmp_path / "preview.docx"
    doc.save(str(path))

    with path.open("rb") as handle:
        response = client.post(
            "/api/preview-anonymization",
            files={"file": ("preview.docx", handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    assert response.status_code == 200
    body = response.json()
    assert "договор" in body["original_text"].lower()
    assert body["input_type"] == "file"


def test_archive_list(client: TestClient) -> None:
    response = client.get("/api/archive")
    assert response.status_code == 200
    assert "cases" in response.json()


def test_preview_anonymization_text(client: TestClient) -> None:
    text = (
        "Иванов Иван Иванович заключил договор. "
        "ИНН 771234567890. Тел: +7 999 000 11 22. "
    ) * 3
    response = client.post("/api/preview-anonymization", data={"text": text})
    assert response.status_code == 200
    body = response.json()
    assert "anonymized_text" in body
    assert "replacements" in body
    assert isinstance(body["replacements"], list)
    assert len(body["replacements"]) > 0
    record = body["replacements"][0]
    assert all(key in record for key in ["id", "original", "placeholder", "entity_type", "source"])
    assert body["total_replacements"] > 0
    assert "review_candidates" in body


def test_preview_no_input_fails(client: TestClient) -> None:
    response = client.post("/api/preview-anonymization", data={})
    assert response.status_code == 422


def test_legal_terms_whitelist() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Работник обратился в Верховный суд РФ согласно ТК РФ статье 81."
    result = anonymizer.anonymize(text)
    assert "Верховный суд" in result.anonymized_text
    assert "ТК РФ" in result.anonymized_text


def test_replacement_records_populated() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Петров Петр Петрович, ИНН 500100732259, тел +7 (495) 000-00-00"
    result = anonymizer.anonymize(text)
    assert len(result.replacements) > 0
    for record in result.replacements:
        assert record.id
        assert record.original
        assert record.placeholder
        assert 0.0 <= record.confidence <= 1.0


def test_process_with_pre_anonymized(client: TestClient) -> None:
    pre_text = "Договор заключён между [ФИО] и [ОРГАНИЗАЦИЯ]. Предмет: аренда."
    response = client.post(
        "/api/process-text",
        json={
            "text": "оригинальный текст не важен здесь",
            "pre_anonymized_text": pre_text,
            "confirmed_replacements": 2,
            "confirmed_entities": {"PER": ["кто-то"]},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["anonymization"]["total_replacements"] == 2


def test_lifespan_no_deprecation_warning() -> None:
    from pathlib import Path

    source = Path("backend/main.py").read_text(encoding="utf-8")
    assert "on_event" not in source, "on_event found in main.py — replace with lifespan"


def test_live_http_roundtrip(
    tmp_path: pytest.TempPathFactory,
    monkeypatch: pytest.MonkeyPatch,
    free_tcp_port_factory: Callable[[], int],
) -> None:
    """Run a real HTTP roundtrip against a live uvicorn server."""

    monkeypatch.setenv("KP_API_KEY", "DEMO_KEY")
    monkeypatch.delenv("LLM_BASE_URL", raising=False)
    monkeypatch.delenv("LLM_MODEL", raising=False)
    monkeypatch.setenv("ANALYSIS_MODE", "auto")

    from backend.core import archive as archive_module
    from backend.core.archive import ArchiveManager
    import backend.main as main_module
    import uvicorn

    temp_root = tmp_path
    uploads_dir = temp_root / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    temp_archive = ArchiveManager(temp_root / "live_cases.db")

    monkeypatch.setattr(main_module, "archive", temp_archive)
    monkeypatch.setattr(archive_module, "archive", temp_archive)
    monkeypatch.setattr(main_module, "UPLOAD_DIR", uploads_dir)

    port = free_tcp_port_factory()
    config = uvicorn.Config(
        main_module.app,
        host="127.0.0.1",
        port=port,
        log_level="warning",
    )
    server = uvicorn.Server(config)
    main_module.app.state.shutdown_handler = lambda: setattr(server, "should_exit", True)
    server_thread = threading.Thread(target=server.run, daemon=True)
    server_thread.start()

    try:
        with httpx.Client(base_url=f"http://127.0.0.1:{port}", timeout=5.0) as client:
            health_response = None
            for _ in range(40):
                try:
                    health_response = client.get("/health")
                    if health_response.status_code == 200:
                        break
                except httpx.HTTPError:
                    pass
                time.sleep(0.25)

            assert health_response is not None
            assert health_response.status_code == 200
            assert health_response.json()["status"] == "ok"

            process_response = client.post(
                "/api/process-text",
                json={
                    "text": (
                        "Договор аренды заключён между Ивановым Иваном и ООО Ромашка. "
                        "ИНН 7712345678. Телефон +7 999 123 45 67. Адрес: г. Москва, ул. Ленина, д.1"
                    )
                },
            )
            assert process_response.status_code == 200
            process_body = process_response.json()
            assert process_body["anonymization"]["total_replacements"] >= 1

            archive_response = client.get("/api/archive")
            assert archive_response.status_code == 200
            archive_body = archive_response.json()
            assert archive_body["cases"]

            case_id = process_body["case_id"]
            case_response = client.get(f"/api/archive/{case_id}")
            assert case_response.status_code == 200
            case_body = case_response.json()
            assert case_body["input_type"] == "manual"
            assert case_body["filename"]

            shutdown_response = client.post("/api/system/shutdown")
            assert shutdown_response.status_code == 200
            assert shutdown_response.json()["status"] == "stopping"

            stopped = False
            for _ in range(40):
                try:
                    client.get("/health")
                except httpx.HTTPError:
                    stopped = True
                    break
                time.sleep(0.1)
            assert stopped is True
    finally:
        server.should_exit = True
        server_thread.join(timeout=10)


def test_registry_demo_mode() -> None:
    from backend.adapters.consultant_plus import StubAdapter
    from backend.adapters.registry import AnalysisMode, registry

    adapter = registry.get(AnalysisMode.DEMO)
    assert isinstance(adapter, StubAdapter)


def test_registry_auto_returns_stub_when_nothing_configured() -> None:
    import os
    from unittest.mock import patch

    from backend.adapters.consultant_plus import StubAdapter
    from backend.adapters.registry import AdapterRegistry, AnalysisMode

    with patch.dict(os.environ, {"KP_API_KEY": "", "LLM_BASE_URL": "", "ANALYSIS_MODE": "auto"}, clear=False):
        adapter = AdapterRegistry().get(AnalysisMode.AUTO)
        assert isinstance(adapter, StubAdapter)


def test_llm_prompt_builder() -> None:
    from backend.adapters.local_llm import LegalPromptBuilder

    builder = LegalPromptBuilder()
    system, user = builder.build("Договор аренды между сторонами.")
    assert "российский юрист" in system.lower()
    assert "Договор аренды" in user
    assert "JSON" in user
    assert "summary" in user
    assert "relevant_laws" in user


def test_llm_parse_valid_response() -> None:
    from backend.adapters.local_llm import LocalLLMAdapter

    adapter = LocalLLMAdapter()
    raw = """{"summary": "Договор аренды",
              "legal_area": "гражданское",
              "relevant_laws": [{"title": "ГК РФ", "article": "ст.606", "text": "..."}],
              "court_practice": [],
              "recommendations": "Проверить срок аренды",
              "risks": "Отсутствует акт приёма-передачи",
              "confidence": 0.88}"""
    result = adapter._parse_llm_response(raw)
    assert result["summary"] == "Договор аренды"
    assert result["confidence"] == 0.88
    assert len(result["relevant_laws"]) == 1


def test_llm_parse_with_markdown_fences() -> None:
    from backend.adapters.local_llm import LocalLLMAdapter

    adapter = LocalLLMAdapter()
    raw = (
        "```json\n"
        '{"summary": "тест", "relevant_laws": [], "court_practice": [], "recommendations": "r", "confidence": 0.7}\n'
        "```"
    )
    result = adapter._parse_llm_response(raw)
    assert result["summary"] == "тест"


def test_process_text_with_demo_mode(client: TestClient) -> None:
    text = "Договор аренды заключён между сторонами. " * 5
    response = client.post(
        "/api/process-text",
        json={
            "text": text,
            "analysis_mode": "demo",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["legal_analysis"]["source"] is not None


def test_health_returns_adapter_status(client: TestClient) -> None:
    response = client.get("/health")
    assert response.status_code == 200
    body = response.json()
    assert "kp_available" in body
    assert "llm_available" in body
    assert body["version"] == "4.0.1"


def test_kp_test_endpoint_with_demo_key(client: TestClient) -> None:
    response = client.post("/api/kp/test", json={"api_key": "DEMO_KEY"})
    assert response.status_code == 200
    assert response.json()["ok"] is False


def test_compare_cases_endpoint(client: TestClient) -> None:
    left = client.post(
        "/api/process-text",
        json={"text": ("Договор аренды между ООО Альфа и ИП Петров. " * 5), "analysis_mode": "demo"},
    )
    right = client.post(
        "/api/process-text",
        json={"text": ("Трудовой спор о восстановлении на работе. " * 5), "analysis_mode": "demo"},
    )

    response = client.post(
        "/api/compare",
        json={"left_case_id": left.json()["case_id"], "right_case_id": right.json()["case_id"]},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["left_case"]["id"] == left.json()["case_id"]
    assert body["right_case"]["id"] == right.json()["case_id"]
    assert 0.0 <= body["document"]["similarity"] <= 1.0
    assert "analysis" in body


def test_export_docx_endpoint(client: TestClient) -> None:
    response = client.post(
        "/api/process-text",
        json={"text": ("Договор поставки между сторонами. " * 5), "analysis_mode": "demo"},
    )
    case_id = response.json()["case_id"]

    export_response = client.get(f"/api/archive/{case_id}/export?format=docx")
    assert export_response.status_code == 200
    assert export_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert export_response.content[:2] == b"PK"


def test_export_pdf_endpoint(client: TestClient) -> None:
    response = client.post(
        "/api/process-text",
        json={"text": ("Исковое заявление о взыскании задолженности. " * 5), "analysis_mode": "demo"},
    )
    case_id = response.json()["case_id"]

    export_response = client.get(f"/api/archive/{case_id}/export?format=pdf")
    assert export_response.status_code == 200
    assert export_response.headers["content-type"].startswith("application/pdf")
    assert export_response.content.startswith(b"%PDF")


def test_audit_log_records_review_and_open(client: TestClient) -> None:
    process_response = client.post(
        "/api/process-text",
        json={"text": ("Договор аренды между сторонами. " * 5), "analysis_mode": "demo"},
    )
    case_id = process_response.json()["case_id"]

    open_response = client.get(f"/api/archive/{case_id}")
    review_response = client.patch(f"/api/archive/{case_id}/review")

    assert open_response.status_code == 200
    assert review_response.status_code == 200

    audit_response = client.get("/api/audit")
    assert audit_response.status_code == 200
    actions = {entry["action"] for entry in audit_response.json()["entries"]}
    assert "analysis_completed" in actions
    assert "case_opened" in actions
    assert "case_reviewed" in actions


def test_audit_log_filter_by_case(client: TestClient) -> None:
    first = client.post(
        "/api/process-text",
        json={"text": ("Первый договор аренды. " * 8), "analysis_mode": "demo"},
    ).json()["case_id"]
    second = client.post(
        "/api/process-text",
        json={"text": ("Второй договор подряда. " * 8), "analysis_mode": "demo"},
    ).json()["case_id"]

    _ = client.get(f"/api/archive/{first}")
    _ = client.get(f"/api/archive/{second}")

    audit_response = client.get(f"/api/audit?case_id={first}")
    assert audit_response.status_code == 200
    entries = audit_response.json()["entries"]
    assert entries
    assert all(entry["case_id"] == first for entry in entries)


def test_export_and_compare_write_audit_log(client: TestClient) -> None:
    left = client.post(
        "/api/process-text",
        json={"text": ("Договор лизинга между сторонами. " * 6), "analysis_mode": "demo"},
    ).json()["case_id"]
    right = client.post(
        "/api/process-text",
        json={"text": ("Дополнительное соглашение к договору лизинга. " * 6), "analysis_mode": "demo"},
    ).json()["case_id"]

    _ = client.get(f"/api/archive/{left}/export?format=docx")
    _ = client.get(f"/api/archive/{left}/export?format=pdf")
    _ = client.post("/api/compare", json={"left_case_id": left, "right_case_id": right})

    audit_response = client.get(f"/api/audit?case_id={left}")
    assert audit_response.status_code == 200
    actions = {entry["action"] for entry in audit_response.json()["entries"]}
    assert "export_docx" in actions
    assert "export_pdf" in actions
    assert "cases_compared" in actions


def test_anonymizer_rejects_tabular_header_sequences_as_persons() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Оплачено Сумма Вид\nПериод Сумма Дни"
    result = anonymizer.anonymize(text)

    assert result.anonymized_text == text
    assert not any(record.entity_type == "PER" for record in result.replacements)


def test_anonymizer_rejects_inflected_contract_roles() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Права продавца, покупателя и стороны договора согласованы."
    result = anonymizer.anonymize(text)

    assert result.anonymized_text == text
    assert not any(record.entity_type in {"PER", "ORG"} for record in result.replacements)


def test_anonymizer_rejects_generic_management_word_as_org() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Решение управления по вопросу стороны договора исполнено."
    result = anonymizer.anonymize(text)

    assert result.anonymized_text == text
    assert not any(record.entity_type == "ORG" for record in result.replacements)


def test_anonymizer_strips_person_role_prefix_from_full_name() -> None:
    from backend.core.anonymizer import anonymizer

    text = "Гражданин Иванов Иван Иванович подал заявление в суд."
    result = anonymizer.anonymize(text)

    assert "[ФИО1]" in result.anonymized_text
    per_records = [record for record in result.replacements if record.entity_type == "PER"]
    assert len(per_records) == 1
    assert per_records[0].original == "Иванов Иван Иванович"
    assert "гражданин" not in per_records[0].original.casefold()


def test_person_aliases_ignore_leading_noise_word() -> None:
    from backend.core.anonymizer import anonymizer

    text = (
        "Иванов Иван Иванович подписал договор. "
        "Далее Иванов направил копию в суд."
    )
    result = anonymizer.anonymize(text)

    assert result.anonymized_text.count("[ФИО1]") == 2
    assert "[ФИО2]" not in result.anonymized_text


def test_org_rule_marks_borderline_org_for_review() -> None:
    from backend.core.anonymizer import _CandidateSpan
    from backend.core.entity_rules.org_rules import OrganizationRuleEngine

    text = "Вектор Гранит направило письмо."
    candidate = _CandidateSpan(
        id="org-review-1",
        original="Вектор Гранит",
        placeholder="[ОРГАНИЗАЦИЯ]",
        entity_type="ORG",
        start=0,
        end=len("Вектор Гранит"),
        source="ner",
        confidence=0.7,
    )

    decision = OrganizationRuleEngine("[ОРГАНИЗАЦИЯ]").validate_candidate(text, candidate)
    assert decision.reviewable is True
    assert decision.accepted is False


def test_preview_exposes_review_candidates(client: TestClient) -> None:
    response = client.post("/api/preview-anonymization", data={"text": "Вектор Гранит направило письмо."})
    assert response.status_code == 200
    body = response.json()
    assert "review_candidates" in body
    assert isinstance(body["review_candidates"], list)


def test_anonymizer_case_corpus_metrics() -> None:
    from backend.core.anonymizer import anonymizer
    from backend.core.entity_rules.common import canonicalize_org_name, clean_entity_text, normalize_entity_text

    fixture_path = Path("tests/fixtures/anonymizer_cases.jsonl")
    rows = [json.loads(line) for line in fixture_path.read_text(encoding="utf-8-sig").splitlines() if line.strip()]
    assert len(rows) >= 215

    metrics = {
        "PER": {"tp": 0, "fp": 0, "fn": 0},
        "ORG": {"tp": 0, "fp": 0, "fn": 0},
        "АДРЕС": {"tp": 0, "fp": 0, "fn": 0},
    }

    def _addr_key(value: str) -> str:
        return normalize_entity_text(clean_entity_text(value))

    for row in rows:
        result = anonymizer.anonymize(row["text"])
        found = {
            "PER": {
                normalize_entity_text(clean_entity_text(record.original))
                for record in result.replacements
                if record.entity_type == "PER"
            },
            "ORG": {
                canonicalize_org_name(record.original)
                for record in result.replacements
                if record.entity_type == "ORG"
            },
            "АДРЕС": {
                _addr_key(record.original)
                for record in result.replacements
                if record.entity_type == "АДРЕС"
            },
        }

        for entity_type in ("PER", "ORG", "АДРЕС"):
            raw_expected = row["expected"].get(entity_type, [])
            if entity_type == "PER":
                expected = {normalize_entity_text(clean_entity_text(item)) for item in raw_expected}
            elif entity_type == "ORG":
                expected = {canonicalize_org_name(item) for item in raw_expected}
            else:
                expected = {_addr_key(item) for item in raw_expected}
            # Address matching is substring-lenient — a detected address may
            # include slightly more context than the fixture snapshot.
            if entity_type == "АДРЕС":
                matched_expected = set()
                for exp in expected:
                    if any(exp in det or det in exp for det in found[entity_type]):
                        matched_expected.add(exp)
                matched_found = set()
                for det in found[entity_type]:
                    if any(exp in det or det in exp for exp in expected):
                        matched_found.add(det)
                metrics[entity_type]["tp"] += len(matched_expected)
                metrics[entity_type]["fp"] += len(found[entity_type] - matched_found)
                metrics[entity_type]["fn"] += len(expected - matched_expected)
            else:
                metrics[entity_type]["tp"] += len(found[entity_type] & expected)
                metrics[entity_type]["fp"] += len(found[entity_type] - expected)
                metrics[entity_type]["fn"] += len(expected - found[entity_type])

    per_precision = metrics["PER"]["tp"] / (metrics["PER"]["tp"] + metrics["PER"]["fp"])
    per_recall = metrics["PER"]["tp"] / (metrics["PER"]["tp"] + metrics["PER"]["fn"])
    org_precision = metrics["ORG"]["tp"] / (metrics["ORG"]["tp"] + metrics["ORG"]["fp"])
    org_recall = metrics["ORG"]["tp"] / (metrics["ORG"]["tp"] + metrics["ORG"]["fn"])
    addr_tp = metrics["АДРЕС"]["tp"]
    addr_fp = metrics["АДРЕС"]["fp"]
    addr_fn = metrics["АДРЕС"]["fn"]
    addr_precision = addr_tp / (addr_tp + addr_fp) if (addr_tp + addr_fp) else 1.0
    addr_recall = addr_tp / (addr_tp + addr_fn) if (addr_tp + addr_fn) else 1.0

    assert per_precision >= 0.93, metrics
    assert per_recall >= 0.9, metrics
    assert org_precision >= 0.93, metrics
    assert org_recall >= 0.9, metrics
    assert addr_recall >= 0.8, metrics
    assert addr_precision >= 0.85, metrics
